#!/usr/bin/env python3
"""Generate the AI & Analytics Data Collection Inventory as a formatted DOCX.

Source of truth is AI_Analytics_Data_Collection.md in this directory; this script
renders that content with real Word styling (headings, tables, bold) instead of a
raw markdown-to-docx pass, matching the presentation quality of the other generated
docs in this folder. Re-run after editing the markdown source.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

DOCS_DIR = Path(__file__).resolve().parent
OUTPUT = DOCS_DIR / "AI_Analytics_Data_Collection.docx"

NAVY = RGBColor(0x1F, 0x2A, 0x44)
SLATE = RGBColor(0x3A, 0x3A, 0x3A)
ACCENT = RGBColor(0x4C, 0x6E, 0xF5)


def set_cell_shading(cell, hex_color: str) -> None:
    shd = cell._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color,
    })
    cell._tc.get_or_add_tcPr().append(shd)


def add_runs(paragraph, text: str, base_size: int = 10.5) -> None:
    """Render **bold** and `code` spans from markdown-ish inline text."""
    i = 0
    while i < len(text):
        if text[i:i + 2] == "**":
            end = text.find("**", i + 2)
            if end != -1:
                run = paragraph.add_run(text[i + 2:end])
                run.bold = True
                run.font.size = Pt(base_size)
                i = end + 2
                continue
        if text[i] == "`":
            end = text.find("`", i + 1)
            if end != -1:
                run = paragraph.add_run(text[i + 1:end])
                run.font.name = "Consolas"
                run.font.size = Pt(base_size - 0.5)
                run.font.color.rgb = RGBColor(0xB3, 0x00, 0x60)
                i = end + 1
                continue
        j = i
        while j < len(text) and text[j] not in "*`":
            j += 1
        if j == i:
            j += 1
        run = paragraph.add_run(text[i:j])
        run.font.size = Pt(base_size)
        i = j


def add_heading(doc, text: str, level: int) -> None:
    h = doc.add_heading(level=level)
    h.paragraph_format.space_before = Pt(16 if level <= 2 else 10)
    h.paragraph_format.space_after = Pt(6)
    run = h.add_run(text)
    run.font.color.rgb = NAVY if level <= 2 else ACCENT
    if level == 1:
        run.font.size = Pt(20)
    elif level == 2:
        run.font.size = Pt(15)
    else:
        run.font.size = Pt(12.5)


def add_paragraph(doc, text: str, *, italic: bool = False, style: str | None = None):
    p = doc.add_paragraph(style=style)
    add_runs(p, text)
    if italic:
        for run in p.runs:
            run.italic = True
            run.font.color.rgb = SLATE
    p.paragraph_format.space_after = Pt(8)
    return p


def add_bullet(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    add_runs(p, text)
    p.paragraph_format.space_after = Pt(4)


def add_numbered(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    add_runs(p, text)
    p.paragraph_format.space_after = Pt(4)


def add_table(doc, header: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    table.autofit = True
    hdr_cells = table.rows[0].cells
    for idx, label in enumerate(header):
        hdr_cells[idx].text = ""
        p = hdr_cells[idx].paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        set_cell_shading(hdr_cells[idx], "2F3B63")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = ""
            p = cells[idx].paragraphs[0]
            add_runs(p, value, base_size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def parse_markdown_table(lines: list[str], start: int) -> tuple[list[str], list[list[str]], int]:
    header = [c.strip() for c in lines[start].strip().strip("|").split("|")]
    i = start + 2  # skip header + separator row
    rows: list[list[str]] = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        rows.append(cells)
        i += 1
    return header, rows, i


def build() -> None:
    md_path = DOCS_DIR / "AI_Analytics_Data_Collection.md"
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)

    doc.core_properties.title = "EEC — AI & Analytics Data Collection Inventory"
    doc.core_properties.subject = "What data feeds EEC's AI/ML and analytics features, and why"
    doc.core_properties.author = "YarrowTech / EEC"

    i = 0
    in_blockquote = False
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if not line.strip():
            in_blockquote = False
            i += 1
            continue

        if line.strip() == ">":
            i += 1
            continue
        if line.startswith("> "):
            add_paragraph(doc, line[2:], italic=True)
            i += 1
            continue

        if line.startswith("# "):
            add_heading(doc, line[2:], 1)
            i += 1
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:], 2)
            i += 1
            continue
        if line.startswith("### "):
            add_heading(doc, line[4:], 3)
            i += 1
            continue

        if line.strip() == "---":
            i += 1
            continue

        if line.strip().startswith("|"):
            header, rows, next_i = parse_markdown_table(lines, i)
            add_table(doc, header, rows)
            i = next_i
            continue

        stripped = line.strip()
        if stripped.startswith("- "):
            add_bullet(doc, stripped[2:])
            i += 1
            continue
        if len(stripped) > 2 and stripped[0].isdigit() and stripped[1:3] in (". ",):
            add_numbered(doc, stripped[3:])
            i += 1
            continue
        # numbered items like "1. text" where number may be >9
        if stripped[:1].isdigit():
            dot = stripped.find(". ")
            if dot != -1 and stripped[:dot].isdigit():
                add_numbered(doc, stripped[dot + 2:])
                i += 1
                continue

        add_paragraph(doc, stripped)
        i += 1

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run(
        "EEC — AI & Analytics Data Collection Inventory | Internal draft — confirm scope, "
        "retention and legal basis before external publication."
    )
    footer_run.italic = True
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = SLATE

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
